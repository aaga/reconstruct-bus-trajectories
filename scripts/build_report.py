"""Build the project report (DOCX) reproducing Huang et al. (ITSC 2023) on
CTA Route 22 SB heartbeat data.

Output: report/Reconstructing_Bus_Trajectories_Report.docx

Designed so the user can open the file in Word/Pages and edit. Figures point
at slides/*.png that already exist in the repo.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
SLIDES = REPO / "slides"
OUT = REPO / "report" / "Reconstructing_Bus_Trajectories_Report.docx"


# --------------------- helpers ---------------------------


def _set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "888888")
        tcBorders.append(e)
    tcPr.append(tcBorders)


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    return h


def add_p(doc, text, italic=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = italic
    return p


def add_runs(doc, runs, size=10.5, alignment=None):
    """runs: list of (text, italic, bold) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if alignment is not None:
        p.alignment = alignment
    for txt, ital, bold in runs:
        r = p.add_run(txt)
        r.font.size = Pt(size)
        r.font.italic = ital
        r.font.bold = bold
    return p


def add_p_subs(doc, parts, size=10.5):
    """Paragraph with inline subscripts. parts: list of (text, role) where
    role is 'normal' or 'sub'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    for txt, role in parts:
        r = p.add_run(txt)
        r.font.size = Pt(size)
        if role == "sub":
            r.font.subscript = True
    return p


def add_fig(doc, path, width_in, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.italic = True
    cap.paragraph_format.space_after = Pt(8)


def add_table(doc, header, rows, widths_in=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = t.rows[0].cells
    for j, h in enumerate(header):
        hdr[j].text = ""
        p = hdr[j].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        _set_cell_border(hdr[j])
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            _set_cell_border(cell)
    if widths_in:
        for j, w in enumerate(widths_in):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    return t


# --------------------- document --------------------------


def build() -> Path:
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()

    # Page setup: 1" margins, 10.5pt body
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ---------- Title ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Reconstructing CTA Route 22 Bus Trajectories from Heartbeat Data")
    r.font.size = Pt(16); r.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A reproduction and extension of Huang et al. (ITSC 2023)")
    r.font.size = Pt(11); r.font.italic = True

    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = auth.add_run("Ashwin Agarwal · MIT")
    r.font.size = Pt(10); r.font.italic = False
    auth.paragraph_format.space_after = Pt(8)

    # ---------- Abstract ----------
    p = doc.add_paragraph()
    r = p.add_run("Abstract  —  ")
    r.bold = True; r.font.size = Pt(10)
    r = p.add_run(
        "We reproduce the LOCREG-PCHIP trajectory-reconstruction algorithm of "
        "Huang et al. on a different agency (CTA), a different corridor "
        "(Route 22 Clark, southbound), and a substantially lower-cadence GPS "
        "feed (~30 s vs. the paper's ~6 s heartbeat). To make the experiment "
        "feasible at scale we (a) built an open-source, multi-agency real-time "
        "scraper that compacts vehicle pings into Hive-partitioned Parquet on "
        "Cloudflare R2, and (b) replaced the paper's Valhalla-per-ping "
        "map-matching with direct projection onto the GTFS shape polyline. "
        "Across nine days we reconstructed 431 complete southbound trips. "
        "We then extend the paper with an OSM-driven intersection enrichment "
        "step and a simple two-tier delay-attribution heuristic that assigns "
        "every <5 mph interval to a nearby bus stop or controlled "
        "intersection. The reproduced trajectories qualitatively match the "
        "paper's findings (zero speed at stops, smooth deceleration into "
        "signals); the attribution layer surfaces Devon, Foster and Addison "
        "as the dominant delay sources on Clark Street, with traffic signals "
        "near West Adams and West Monroe contributing a comparable share "
        "downtown."
    )
    r.font.size = Pt(10); r.font.italic = False
    p.paragraph_format.space_after = Pt(10)

    # ---------- 1. Problem ----------
    add_h(doc, "1.  Problem statement", level=1)
    add_p(doc,
          "Transit agencies are increasingly exposing high-resolution GPS "
          "(\"heartbeat\") feeds for their fleets. Although these feeds are far "
          "richer than stop-level AVL, the data are noisy in space (vehicles "
          "appear to drift off-road) and irregular in time (cadence varies "
          "from sub-second to minutes). Huang et al. show that, given only "
          "(latitude, longitude, timestamp), one can reconstruct a "
          "continuous, smooth, monotone vehicle trajectory f(t) from which "
          "speed v(t)=f'(t) and acceleration a(t)=f''(t) are recovered "
          "analytically. We adopt their framing and ask: does the same "
          "algorithm work on a feed whose median cadence is ~30 s rather "
          "than 6 s, on a different agency's API, and at scale across "
          "hundreds of trips? And once we have many reconstructed "
          "trajectories, can we attribute the resulting delay to specific "
          "infrastructure (stops, signals)?"
          )

    # ---------- 2. Mathematical formulation ----------
    add_h(doc, "2.  Formulation", level=1)
    add_runs(doc, [
        ("We follow Huang et al.'s notation. Let a single trip be a sequence "
         "of timestamped GPS coordinates "
         "(", False, False),
        ("Sᵢ", True, False),
        (", ", False, False),
        ("Cᵢ", True, False),
        (")", False, False),
        (" for i = 1,…,n, where ", False, False),
        ("Cᵢ = (latᵢ, lonᵢ)", True, False),
        (". Map matching projects each raw point onto a fixed route geometry "
         "and produces a time-into-trip / distance-into-trip pair "
         "(", False, False),
        ("tᵢ, dᵢ", True, False),
        (") with t₁ = 0, d₁ = 0. The smoothing step recovers a function "
         "f : ℝ → ℝ that maps time-into-trip to distance, treating each "
         "observed ", False, False),
        ("dᵢ", True, False),
        (" as a noisy realisation ", False, False),
        ("dᵢ = xᵢ + εᵢ", True, False),
        (" of a latent true position. Speed and acceleration follow as "
         "v(t) = f′(t), a(t) = f″(t).", False, False),
    ])

    add_p(doc,
          "An ideal reconstruction (per the paper) is monotone non-decreasing "
          "(buses do not move backwards), at least once differentiable (so "
          "v(t) is continuous), and built from cubic pieces (the lowest "
          "polynomial degree that can encode bus kinematics). The paper "
          "compares four candidate smoothers — linear (LSEG), PCHIP, LOCREG, "
          "and the hybrid LOCREG-PCHIP — and selects LOCREG-PCHIP. We "
          "reproduce LOCREG-PCHIP only.")

    add_h(doc, "2.1  Map matching: shape projection", level=2)
    add_p(doc,
          "Huang et al. send every raw ping to the Valhalla \"trace_attributes\" "
          "service, then convert the matched (road segment, fractional "
          "position) tuple into a distance-along-route value. We initially "
          "stood up a Valhalla instance with an OSM extract of Cook County "
          "and verified it works on a sample trip, but at the corridor scale "
          "we are interested in (~10⁵ pings) the round-trip latency is "
          "prohibitive and Valhalla occasionally mis-snaps to parallel side "
          "streets when GPS noise is large. Because every CTA Route 22 trip "
          "follows a known GTFS shape (id 67803936), we collapse the two-step "
          "snap-then-test into a single projection: precompute the shape "
          "polyline once, and for each ping return the nearest point on the "
          "polyline along with its cumulative distance from the start "
          "(SnapToShapeMatcher in src/bus_trajectories/mapmatch/shape_snap.py). "
          "This costs O(|polyline|) per ping, runs in pure NumPy, and "
          "deterministically constrains the bus to its scheduled route. "
          "It does mean we cannot detect off-route running (deadheading, "
          "deviations); for an in-service trip this is rarely a problem.")

    add_h(doc, "2.2  LOCREG-PCHIP smoothing", level=2)
    add_runs(doc, [
        ("At each ping i we fit a degree-p polynomial in centred time "
         "x = t − tᵢ to the k nearest neighbours (in time), weighted by the "
         "tricube kernel ", False, False),
        ("wₖ = (1 − |xₖ/h|³)³", True, False),
        (" with bandwidth h = max|xₖ|. The smoothed value is "
         "x̃ᵢ ≔ p(0) = â₀, the constant term of the fit. Higher-order "
         "coefficients are discarded. After applying LOCREG to every ping we "
         "forward-fill any monotonicity violation (xᵢ ≔ max(xᵢ, xᵢ₋₁)) and "
         "interpolate the cleaned sequence with a PCHIP spline. The result "
         "is monotone, C¹, and made of cubic pieces — meeting the paper's "
         "three ideal properties.", False, False),
    ])
    add_p(doc,
          "We deviate from the paper on the bandwidth k. The paper uses "
          "k = 20 with a 6 s median cadence, so each window spans roughly "
          "two minutes of trip time. CTA's BusTime API publishes vehicle "
          "positions every ~30 s; running k = 20 on our data would smooth "
          "across ten minutes of trip and erase real stops and signals. We "
          "therefore adopt k = 5, which spans ~2.5 min and is close to the "
          "paper's effective time window. A consequence we report explicitly "
          "in the appendix is that with k = 5 the tricube zeroes the two "
          "boundary points exactly, leaving 3–4 effective rows for 4 "
          "unknowns; LOCREG often interpolates rather than smooths the "
          "central pings, and most of the visible smoothness in our "
          "time-space diagrams comes from the PCHIP step.")

    add_h(doc, "2.3  Intersection enrichment (extension)", level=2)
    add_p(doc,
          "The paper notes that detecting bus interactions with infrastructure "
          "(traffic signals, pedestrian crossings) is enabled by the "
          "trajectory but does not provide a method. We add a small offline "
          "step: for every OSM way the GTFS shape traverses, fetch all "
          "referenced nodes via Overpass and keep those (i) belonging to ≥ 2 "
          "highway-tagged ways (i.e., true intersections), (ii) projecting "
          "to within 30 m of the shape, and (iii) controlled in the bus's "
          "direction of travel — meaning either tagged "
          "highway=traffic_signals (signals control all approaches) or "
          "carrying a stop / give_way sign on the bus's approach with a "
          "compatible direction tag. Output is a per-shape JSON of "
          "ControlPoints with intersection node id, latitude/longitude, "
          "distance along the route, control type, and cross-street name. "
          "Stop signs and signals on side streets the bus crosses freely are "
          "correctly excluded.")

    add_h(doc, "2.4  Delay attribution (extension)", level=2)
    add_p_subs(doc, [
        ("Given a reconstructed trajectory, we segment it into \"slow "
         "windows\" where v(t) < 5 mph for at least 2 s. For each window "
         "[t", "normal"), ("a", "sub"),
        (", t", "normal"), ("b", "sub"),
        ("] with corresponding distance interval [x", "normal"), ("a", "sub"),
        (", x", "normal"), ("b", "sub"),
        ("] and duration Δ = t", "normal"), ("b", "sub"),
        (" − t", "normal"), ("a", "sub"),
        (":", "normal"),
    ])
    add_p_subs(doc, [
        ("    (i)  if any GTFS bus stops fall in [x", "normal"), ("a", "sub"),
        (", x", "normal"), ("b", "sub"),
        ("], distribute Δ evenly across them;", "normal"),
    ])
    add_p_subs(doc, [
        ("    (ii) else, if any controlled intersections fall in "
         "[x", "normal"), ("a", "sub"),
        (", x", "normal"), ("b", "sub"),
        (" + 0.05 mi] (a direction-aware lookahead — buses "
         "decelerate before reaching a red light), distribute Δ evenly "
         "across them;", "normal"),
    ])
    add_p(doc,
          "    (iii) else, attribute Δ to a residual OTHER bucket.")
    add_p(doc,
          "Per-feature totals are summed over all 431 trips to produce "
          "ranked total / mean / std / quartile statistics. The attribution "
          "is intentionally heuristic — it ignores partial spatial overlap, "
          "splits delay equally when several features are active, and cannot "
          "explain the OTHER bucket — but it gives a consistent first-cut "
          "summary of where on the corridor delay accumulates.")

    add_h(doc, "3.  Data and pipeline", level=1)
    add_p(doc,
          "We built scrape-bus-pings, an open-source long-running scraper "
          "that polls four agencies (MBTA, MTA NYC Bus, TfL, CTA) every "
          "15 seconds, batches the canonicalised vehicle positions into "
          "1-minute Parquet files, ships them to Cloudflare R2, and an "
          "in-process compactor merges each completed UTC hour into a single "
          "Hive-partitioned Parquet object. Files are read lazily over the "
          "public R2 URL with DuckDB or pyarrow. For this study we use the "
          "CTA BusTime feed (canonicalised to the same 26-column GTFS-RT "
          "schema as the other agencies) collected over nine days "
          "(2026-04-27 to 2026-05-06), which produced 217 hour-files. "
          "Filtering to Route 22 with a strict south-bound classifier "
          "(see §3.1) yields 431 complete trips on pattern 3936.")

    add_p(doc,
          "Trip selection rules: (i) ≥ 30 raw pings; (ii) the first ping "
          "snaps within 600 m of the shape origin (excludes mid-route block "
          "starts); (iii) the trajectory eventually reaches within 50 m of "
          "the shape terminus (we then truncate the tail so post-arrival "
          "layover does not pollute the smoother); (iv) no inter-ping gap "
          "exceeds 5 min in the truncated series. Trip-id collisions across "
          "days are resolved by prefixing the date to form a unique key. "
          "Of 1734 candidate trips, 431 pass all four filters.")

    # Pipeline figure
    add_fig(doc, SLIDES / "C_pipeline.png", 6.5,
            "Figure 1.  Reconstruction pipeline. Raw heartbeat pings from R2 "
            "are projected onto the GTFS shape polyline (one-step "
            "shape-snap; deviation from Huang et al.'s Valhalla-per-ping), "
            "then passed to LOCREG-PCHIP smoothing. The resulting f(t) is "
            "differentiated to produce v(t) and a(t), and overlaid against "
            "OSM-derived intersections and GTFS stops for delay attribution.")

    add_h(doc, "4.  Results", level=1)
    add_p(doc,
          "Applied to the 431-trip population, the LOCREG-PCHIP pipeline "
          "produces a coherent envelope of trajectories that differ only in "
          "their realised speeds (Figure 2). The fastest trips traverse the "
          "10.7 mi route in ~50 min; the slowest take ≥ 90 min, primarily "
          "during the PM peak. No per-trip tuning was needed.")

    # Multi-trip aligned (now in Results)
    add_fig(doc, SLIDES / "F4_timespace_alltrips_aligned.png", 6.0,
            "Figure 2.  Time-space diagrams of all 431 reconstructed Route 22 "
            "SB trips, aligned to the bus's actual departure (t = 0 = last "
            "stationary ping before first ≥ 0.03 mi forward progress). The "
            "envelope's lower edge is the fastest observed traversal "
            "(~50 min for 10.7 mi); the upper edge shows trips taking "
            "≥ 90 min, primarily during the PM peak.")

    add_p(doc,
          "On a randomly chosen single trip (CTA trip_id 1001350, "
          "2026-04-29 PM peak), the reconstructed trajectory exhibits all "
          "four qualitative behaviours that Huang et al. report on MBTA "
          "Route 1: (a) zero speed at every served bus stop; (b) brief, "
          "smooth decelerations to ~0 mph at red signals on Clark; "
          "(c) speeds of 25–30 mph on the relatively unobstructed stretch "
          "around Devon; and (d) extended dwell at the southern terminal as "
          "the bus enters layover. Without access to CTA's door-open AVL "
          "stream we cannot reproduce the paper's quantitative validation "
          "(percentage of door-open seconds at zero speed); we instead rely "
          "on visual inspection and the high-cadence speed profile shown in "
          "Figure 3.")

    # Speed profile w/ shading
    add_fig(doc, SLIDES / "B4_speed.png", 6.0,
            "Figure 3.  Speed profile of trip 1001350 between miles 3.5 and "
            "5.5 of the route, with bus stops (blue), traffic signals "
            "(amber dashed), and v < 5 mph delay segments (red shading) "
            "overlaid. Each speed dip aligns either with a stop, a signal, "
            "or both — qualitatively matching Huang et al.'s Figure 6, "
            "though their MBTA feed gives a smoother profile thanks to "
            "5–10× higher cadence.")

    add_p(doc,
          "Aggregating across the 431-trip population, the delay-attribution "
          "heuristic charges a mean of 26.3 min/trip to features with known "
          "identity and 5.9 min/trip to OTHER (mostly congested mid-block "
          "running and bunching). Figure D1 in Appendix D ranks the top 25 "
          "delay sources by mean attributed delay per trip: bus stops at "
          "Devon (0.92), Foster (0.92), Addison (0.90), Belmont (0.79), and "
          "Diversey (0.77) lead, while the heaviest signal-driven delays "
          "occur at West Adams (0.54) and West Monroe (0.56) downtown. "
          "Figure D2 plots the same data along the route; the spatial "
          "concentration in Edgewater/Lakeview (miles 1.5–6) and downtown "
          "(miles 9–10.5) is evident.")

    add_h(doc, "5.  Discussion, limitations and future work", level=1)
    add_p(doc,
          "Reproducibility. The core LOCREG-PCHIP algorithm transfers "
          "directly from MBTA Route 1 (6 s cadence) to CTA Route 22 (30 s "
          "cadence) once the bandwidth is rescaled. The rescaling is "
          "non-trivial: with k = 5 the tricube zeros the boundary "
          "neighbours and the smoother degenerates to a near-interpolant "
          "(see Appendix A.1). At MBTA's cadence one could plausibly use "
          "k = 20 unmodified; at CTA's cadence we must accept either a "
          "near-interpolating smoother or a window so wide it averages "
          "across stops. We chose the former.")
    add_p(doc,
          "Map-matching trade-off. Replacing Valhalla with shape-snap is a "
          "factor of ~10² speed-up and removes a network dependency, but "
          "permanently tethers each trip to its scheduled shape. Genuine "
          "off-route events (detours, emergency reroutes, mis-tagged "
          "patterns) become invisible. For the 22-Clark corridor, on which "
          "essentially all SB trips run pattern 3936, the trade-off is "
          "favourable; on routes with significant pattern variation it "
          "would not be.")
    add_p(doc,
          "Validation gap. The paper's strongest validation is "
          "speed-at-door-open events from MBTA AVL. CTA's BusTime API does "
          "not expose door state, so we cannot reproduce that. A natural "
          "next step is to combine the trajectory with any of the agency's "
          "internal AVL extracts (which include door-open events) or with "
          "automatic passenger counter (APC) data when available.")
    add_p(doc,
          "Attribution heuristic. The two-tier algorithm is deliberately "
          "simple. Its largest weakness is that the OTHER bucket "
          "(~5.9 min/trip on average, ~22% of attributed delay) is a "
          "black box: any congestion not localised near a stop or signal "
          "ends up there. A more principled successor would model delay as "
          "a non-negative sum over feature-specific kernels with a "
          "diffuse-traffic baseline, fit jointly across all trips. We leave "
          "this for follow-on work.")

    # ---------- References ----------
    doc.add_page_break()
    add_h(doc, "References", level=1)
    refs = [
        "Y. Huang, A. Abdelhalim, A. Stewart, J. Zhao, H. Koutsopoulos. "
        "Reconstructing Transit Vehicle Trajectory Using High-Resolution GPS "
        "Data. 2023 IEEE 26th Int. Conf. on Intelligent Transportation "
        "Systems (ITSC), pp. 5247-5253.",
        "F. N. Fritsch and R. E. Carlson. Monotone Piecewise Cubic "
        "Interpolation. SIAM J. Numer. Anal. 17(2):238-246, 1980.",
        "F. N. Fritsch and J. Butland. A Method for Constructing Local "
        "Monotone Piecewise Cubic Interpolants. SIAM J. Sci. Stat. Comput. "
        "5(2):300-304, 1984.",
        "W. S. Cleveland. Robust Locally Weighted Regression and Smoothing "
        "Scatterplots. J. Am. Stat. Assoc. 74(368):829-836, 1979.",
        "T. Toledo, H. N. Koutsopoulos, K. I. Ahmed. Estimation of Vehicle "
        "Trajectories with Locally Weighted Regression. Transp. Res. Rec. "
        "1999:161-169, 2007.",
        "S. P. Venthuruthiyil, M. Chunchu. Trajectory reconstruction using "
        "locally weighted regression. Transportmetrica A 14:1-19, 2018.",
        "Valhalla — open-source routing engine for OpenStreetMap. "
        "https://github.com/valhalla",
        "OpenStreetMap contributors. https://www.openstreetmap.org",
        "Chicago Transit Authority. BusTime developer API. "
        "https://www.transitchicago.com/developers/bustracker/",
        "GTFS Static Specification. https://gtfs.org/schedule/reference/",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        r = p.add_run(f"[{i}]  {ref}"); r.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(2)

    # ---------- Appendix ----------
    doc.add_page_break()
    add_h(doc, "Appendix A.  Worked example: LOCREG step", level=1)
    add_p(doc,
          "To make Section 2.2 concrete, Figure A1 walks through the "
          "weighted-LS fit at one ping of a synthetic 5-point series with "
          "k = 5. The weights w = (0.886, 1.000, 0.897, 0.365, 0.000) zero "
          "out the rightmost neighbour exactly — a structural property of "
          "the tricube at the bandwidth boundary. The resulting cubic "
          "p(x) = −2.98 × 10⁻⁴ x³ + 1.24 × 10⁻² x² + 5.51 x + 138 evaluated "
          "at x = 0 returns 138 m, exactly the observed dᵢ. Because the "
          "kernel has zeroed two of the five weights at the symmetric centre "
          "of the trip and one at every off-centre interior position, the "
          "system is rank-deficient enough that LOCREG behaves as an "
          "interpolant on the inner pings rather than a true smoother. "
          "The local velocity p′(0) = a₁ ≈ 5.5 m/s ≈ 12 mph is, however, a "
          "useful by-product — it is what one would use if the polynomial "
          "fit's higher coefficients were retained.")

    add_fig(doc, SLIDES / "L_locreg_explainer.png", 6.6,
            "Figure A1.  Worked LOCREG example. Left: the algorithm step by "
            "step, with the numerical example used in the right panel. "
            "Right: data points sized by tricube weight (boundary point at "
            "weight 0 omitted), the fitted local cubic, and the LOCREG "
            "output p(0) = a₀ at the pivot.")

    add_h(doc, "Appendix B.  Intersection enrichment output", level=1)
    add_p(doc,
          "The OSM-driven enrichment (§2.3) produces 47 controlled "
          "intersections on shape 67803936 (Howard → Harrison, 10.7 mi), "
          "all classified as traffic_signals (no give-way nodes, and "
          "stop signs were correctly excluded — Clark Street is an arterial, "
          "so cross streets stop for it, not the reverse). Figure B1 plots "
          "the result on a Leaflet basemap; manual cross-check against "
          "Google Maps confirms every well-known signal on Clark "
          "(Foster, Belmont, Diversey, North Ave, Division, Chicago, "
          "Madison, Harrison) is captured.")

    add_fig(doc, SLIDES / "E_intersections_map.png", 4.4,
            "Figure B1.  Controlled intersections detected by the OSM/Overpass "
            "enrichment for Route 22 SB. Red markers = traffic signals "
            "where the bus does not have free right-of-way.")

    add_h(doc, "Appendix C.  Reconstruction examples at multiple scales", level=1)
    add_fig(doc, SLIDES / "F2_timespace_50trips_aligned.png", 6.4,
            "Figure C1.  Fifty consecutive Route 22 SB trips starting from "
            "07:00 CDT on 2026-05-05, aligned to actual departure. The "
            "fastest trips finish in ~50 min; the slowest take ~90 min, "
            "with the spread driven primarily by the Belmont–Chicago "
            "section.")
    add_fig(doc, SLIDES / "F1_timespace_50trips_clock.png", 6.4,
            "Figure C2.  Same fifty trips, plotted in clock time (07:00–18:00 "
            "UTC). The diagonal banding is the natural service spacing; "
            "wider bands indicate longer headways during the AM/PM peaks.")

    add_h(doc, "Appendix D.  Aggregate delay distribution", level=1)
    add_fig(doc, SLIDES / "H_bar_aggregate.png", 6.4,
            "Figure D1.  Top 25 delay sources across the 431-trip "
            "population, ranked by mean attributed delay per trip. "
            "Whiskers = ± 1 σ. Bus stops dominate the upper half; signals "
            "near the Loop dominate the bottom-right cluster. The "
            "OTHER bucket averages 5.87 ± 4.27 min/trip and is shown only "
            "in the title.")
    add_fig(doc, SLIDES / "H_stem_along_route.png", 6.4,
            "Figure D2.  Same data plotted along the route. Distance along "
            "the shape (mi) on the x-axis; mean delay per trip (min) on "
            "the y-axis. The two delay clusters (Edgewater/Lakeview and "
            "downtown signals) are clearly separated.")

    add_h(doc, "Appendix E.  Implementation summary", level=1)
    add_p(doc,
          "Two repositories support this work. (1) scrape-bus-pings "
          "(~1,400 LOC Python) is a long-running scraper that polls four "
          "agencies' real-time vehicle feeds, canonicalises them to a "
          "shared 26-column GTFS-RT-shaped schema, batches into 1-min "
          "Parquet, uploads to Cloudflare R2, and compacts each completed "
          "UTC hour into a single Hive-partitioned object indexed by a "
          "manifest. The bucket is queryable directly with DuckDB over the "
          "public R2 URL. (2) reconstruct-bus-trajectories (~3,500 LOC "
          "Python) loads the manifest, downloads the relevant hour-files, "
          "filters and reconstructs trajectories, and produces the analysis "
          "and figures used here. The LOCREG-PCHIP smoother is a 165-line "
          "module (src/bus_trajectories/smooth.py); SnapToShapeMatcher is "
          "another 130 lines. All figures in this report were produced by "
          "scripts under scripts/.")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    p = build()
    print(f"saved: {p}  ({p.stat().st_size:,} bytes)")
